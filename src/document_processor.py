import os
import re
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt

from .checklist_verifier import verify_document_checklist
from .report_generator import generate_report
from .comment_inserter import insert_comment
from .rag_engine import get_legal_reference
from .file_utils import read_docx, write_docx

# Patterns for detection
AMBIGUOUS_PATTERNS = [
    r"\bmay\b",
    r"\bendeavour\b",
    r"\bbest endeavours\b",
    r"\bsubject to\b",
    r"\bat its sole discretion\b"
]
JURISDICTION_INDICATORS = [
    "uae federal court",
    "federal court",
    "dubai courts",
    "abu dhabi courts",
    "sharjah court"
]

def _read_doc_text_and_paragraphs(filepath: str):
    return read_docx(filepath)

def _has_signature_block(paragraphs, lookback=8) -> bool:
    if not paragraphs:
        return False
    tail = [p.text.strip().lower() for p in paragraphs[-lookback:] if p.text]
    combined = "\n".join(tail)
    if any(k in combined for k in [
        "signature", "signed", "for and on behalf",
        "authorised signatory", "signature:"
    ]):
        return True
    if re.search(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", combined):
        return True
    return False

def _find_ambiguous_sentences(text: str) -> List[str]:
    if not text:
        return []
    sentences = re.split(r'(?<=[\.\?\!])\s+', text)
    found = []
    for s in sentences:
        low = s.lower()
        for pat in AMBIGUOUS_PATTERNS:
            if re.search(pat, low):
                found.append(s.strip())
                break
    return found

def _has_clause_numbering(paragraphs) -> bool:
    count = 0
    for p in paragraphs:
        t = p.text.strip()
        if re.match(r"^\d+\.", t) or re.match(r"^\d+\.\d+", t) or re.match(r"^\([a-z0-9]+\)", t):
            count += 1
    return count >= 3

def _find_jurisdiction_issues(full_text: str) -> List[Dict[str, Any]]:
    if not full_text:
        return []
    issues = []
    low = full_text.lower()
    for ind in JURISDICTION_INDICATORS:
        if ind in low:
            issues.append({
                "issue": f"Document references '{ind}' which is not ADGM jurisdiction.",
                "severity": "High",
                "suggestion": "Update jurisdiction clause to ADGM Courts.",
                "legal_reference": get_legal_reference("jurisdiction")
            })
    return issues

def _create_front_summary_and_merge(original_doc: Document, issues_by_par: List[Dict[str, Any]], checklist_result: Dict[str, Any]):
    checklist_result = checklist_result or {}
    new_doc = Document()

    # Summary header
    hdr = new_doc.add_paragraph("---- REVIEW SUMMARY (ADGM Corporate Agent - MVP) ----")
    hdr.runs[0].bold = True
    hdr.runs[0].font.size = Pt(12)

    # Checklist info
    new_doc.add_paragraph(f"Detected process: {checklist_result.get('process','Unknown')}")
    new_doc.add_paragraph(f"Documents uploaded: {checklist_result.get('documents_uploaded',0)}")
    new_doc.add_paragraph(f"Required documents: {checklist_result.get('required_documents',0)}")

    missing = checklist_result.get("missing_documents", [])
    if missing:
        new_doc.add_paragraph("Missing documents:")
        for m in missing:
            if isinstance(m, dict):
                new_doc.add_paragraph(f" - {m.get('document')} ({m.get('legal_citation','')})")
            else:
                new_doc.add_paragraph(f" - {m}")

    # Issues section
    new_doc.add_paragraph("")
    new_doc.add_paragraph("Issues found:")
    if not issues_by_par:
        new_doc.add_paragraph(" No issues detected by automated checks.")
    else:
        for idx, it in enumerate(issues_by_par, start=1):
            new_doc.add_paragraph(f"{idx}. [{it.get('severity','')}] {it.get('issue')}")
            if it.get("suggestion"):
                new_doc.add_paragraph(f"   Suggestion: {it['suggestion']}")
            if it.get("legal_reference"):
                new_doc.add_paragraph(f"   Legal Reference: {it['legal_reference']}")
            new_doc.add_paragraph("")

    # Original document after page break
    new_doc.add_page_break()
    for p in original_doc.paragraphs:
        if p.text is None:
            continue
        new_doc.add_paragraph(p.text)
    return new_doc

def add_review_notes_and_save_doc(doc_obj: Document, original_path: str, output_dir: str, issues_by_par: List[Dict[str, Any]], checklist_result: Dict[str, Any]) -> str:
    # Insert inline comments
    for it in issues_by_par:
        para_idx = it.get("paragraph_index")
        comment_text = it.get("issue", "")
        if it.get("suggestion"):
            comment_text += f" Suggestion: {it.get('suggestion')}"
        if it.get("legal_reference"):
            comment_text += f" (Ref: {it.get('legal_reference')})"

        if para_idx is not None and 0 <= para_idx < len(doc_obj.paragraphs):
            insert_comment(doc_obj.paragraphs[para_idx], comment_text)
        else:
            p = doc_obj.add_paragraph()
            insert_comment(p, comment_text)

    # Merge with summary
    merged_doc = _create_front_summary_and_merge(doc_obj, issues_by_par, checklist_result)

    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)
    base = os.path.splitext(os.path.basename(original_path))[0]
    reviewed_name = f"{base}_reviewed.docx"
    reviewed_path = os.path.abspath(os.path.join(output_dir, reviewed_name))

    print(f"[DEBUG] Saving reviewed document to: {reviewed_path}")
    write_docx(merged_doc, reviewed_path)
    print(f"[DEBUG] Successfully saved reviewed document.")

    return reviewed_path

def process_document(filepath: str, checklist_result: Dict[str, Any] = None, output_dir: str = None, debug: bool = False) -> Dict[str, Any]:
    """
    Main document processing function.
    Returns dictionary with file_name, document_type, issues, reviewed_path.
    """
    try:
        doc, paragraphs, full_text = _read_doc_text_and_paragraphs(filepath)
    except Exception as e:
        return {
            "file_name": os.path.basename(filepath),
            "document_type": "Unknown",
            "issues": [{"issue": f"Failed to open .docx: {e}", "severity": "High"}],
            "reviewed_path": None
        }

    # Detect document type
    fn = os.path.basename(filepath).lower()
    if "articles of association" in full_text.lower() or "articles" in fn:
        doc_type = "Articles of Association"
    elif "memorandum of association" in full_text.lower() or "memorandum" in fn:
        doc_type = "Memorandum of Association"
    elif "employment contract" in full_text.lower() or "employment" in fn:
        doc_type = "Employment Contract"
    elif "board resolution" in full_text.lower() or "board resolution" in fn:
        doc_type = "Board Resolution"
    else:
        doc_type = "Unknown"

    # Run checks
    issues = []
    issues.extend(_find_jurisdiction_issues(full_text))

    if not _has_signature_block(paragraphs):
        issues.append({
            "paragraph_index": len(paragraphs)-1 if paragraphs else None,
            "issue": "No signature block detected in the final paragraphs.",
            "severity": "High",
            "suggestion": "Add signature block with name, designation, company name and date.",
            "legal_reference": get_legal_reference("signature")
        })

    ambs = _find_ambiguous_sentences(full_text)
    for s in ambs[:8]:
        idx = next((i for i, p in enumerate(paragraphs) if s in p.text), None)
        issues.append({
            "paragraph_index": idx,
            "issue": f"Ambiguous/non-binding language detected: \"{s[:200]}\"",
            "severity": "Medium",
            "suggestion": "Consider replacing 'may' with 'shall' or more precise wording.",
            "legal_reference": get_legal_reference("ambiguous")
        })

    if not _has_clause_numbering(paragraphs):
        issues.append({
            "paragraph_index": None,
            "issue": "Document appears to lack numbered clauses/section headings.",
            "severity": "Low",
            "suggestion": "Use numbered clause headings (1., 1.1, 2., etc.) to improve clarity.",
            "legal_reference": get_legal_reference("numbered clauses")
        })

    # Save reviewed doc
    reviewed_path = None
    try:
        reviewed_path = add_review_notes_and_save_doc(
            doc,
            filepath,
            output_dir or os.path.join(os.path.dirname(filepath), "tmp_reviewed"),
            issues,
            checklist_result or {}
        )
    except Exception as e:
        print(f"[ERROR] Failed to save reviewed doc for {filepath}: {e}")
        import traceback
        traceback.print_exc()
        reviewed_path = None

    return {
        "file_name": os.path.basename(filepath),
        "document_type": doc_type,
        "issues": issues,
        "reviewed_path": reviewed_path
    }
